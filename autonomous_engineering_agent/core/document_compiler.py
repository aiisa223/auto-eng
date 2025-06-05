"""
Document compilation system for generating technical reports and documentation.
"""

import json
import logging
import os
from datetime import datetime
from typing import Any, Dict, List, Optional, Union

import markdown
from docx import Document
from docx.shared import Inches
from pylatex import Document as LaTeXDocument
from pylatex import Section, Subsection, Command, Package
from pylatex.utils import NoEscape
import time

logger = logging.getLogger(__name__)

class DocumentCompiler:
    """Handles generation of technical documentation in various formats."""
    
    def __init__(self, output_dir: str = "docs"):
        """Initialize the document compiler.
        
        Args:
            output_dir: Directory to store generated documents
        """
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        
    def generate_report(self,
                       content: Dict[str, Any],
                       format: str = "pdf",
                       template: Optional[str] = None) -> str:
        """Generate a technical report.
        
        Args:
            content: Report content and data
            format: Output format (pdf, docx, md)
            template: Optional template to use
            
        Returns:
            Path to the generated document
        """
        if format == "pdf":
            return self._generate_pdf(content, template)
        elif format == "docx":
            return self._generate_docx(content, template)
        elif format == "md":
            return self._generate_markdown(content, template)
        else:
            raise ValueError(f"Unsupported format: {format}")
            
    def _generate_pdf(self,
                     content: Dict[str, Any],
                     template: Optional[str] = None) -> str:
        """Generate a PDF report using LaTeX.
        
        Args:
            content: Report content and data
            template: Optional LaTeX template
            
        Returns:
            Path to the generated PDF
        """
        # Create document
        doc = LaTeXDocument(documentclass="article")
        
        # Add required packages
        doc.packages.append(Package("graphicx"))
        doc.packages.append(Package("amsmath"))
        doc.packages.append(Package("hyperref"))
        
        # Add title
        doc.preamble.append(Command("title", content["title"]))
        doc.preamble.append(Command("author", content["author"]))
        doc.preamble.append(Command("date", NoEscape(r"\today")))
        
        # Add content
        doc.append(Command("maketitle"))
        
        # Add sections
        for section in content["sections"]:
            doc.append(Section(section["title"]))
            doc.append(NoEscape(section["content"]))
            
        # Generate PDF
        output_path = os.path.join(self.output_dir, f"report_{int(time.time())}.pdf")
        doc.generate_pdf(output_path, compiler="pdflatex", clean_tex=True)
        
        return output_path
        
    def _generate_docx(self,
                      content: Dict[str, Any],
                      template: Optional[str] = None) -> str:
        """Generate a Word document.
        
        Args:
            content: Report content and data
            template: Optional Word template
            
        Returns:
            Path to the generated document
        """
        # Create a new Word document
        doc = Document(template) if template else Document()
        
        # Add title
        doc.add_heading(content.get("title", "Technical Report"), 0)
        
        # Add author and date
        doc.add_paragraph(f"Author: {content.get('author', 'Engineering AI')}")
        doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
        
        # Add content
        for section in content.get("sections", []):
            doc.add_heading(section["title"], level=1)
            doc.add_paragraph(section["content"])
            
            # Add subsections
            for subsection in section.get("subsections", []):
                doc.add_heading(subsection["title"], level=2)
                doc.add_paragraph(subsection["content"])
                
        # Add figures
        for figure in content.get("figures", []):
            doc.add_picture(figure["path"], width=Inches(6))
            doc.add_paragraph(figure["caption"])
            
        # Add tables
        for table in content.get("tables", []):
            doc_table = doc.add_table(rows=len(table["data"]), cols=len(table["headers"]))
            doc_table.style = "Table Grid"
            
            # Add headers
            for i, header in enumerate(table["headers"]):
                doc_table.cell(0, i).text = header
                
            # Add data
            for i, row in enumerate(table["data"]):
                for j, cell in enumerate(row):
                    doc_table.cell(i + 1, j).text = str(cell)
                    
        # Save the document
        output_path = os.path.join(
            self.output_dir,
            f"{content.get('title', 'report').lower().replace(' ', '_')}.docx"
        )
        doc.save(output_path)
        
        return output_path
        
    def _generate_markdown(self,
                          content: Dict[str, Any],
                          template: Optional[str] = None) -> str:
        """Generate a Markdown document.
        
        Args:
            content: Report content and data
            template: Optional Markdown template
            
        Returns:
            Path to the generated document
        """
        # Start with template if provided
        md_content = template or ""
        
        # Add title
        md_content += f"# {content.get('title', 'Technical Report')}\n\n"
        
        # Add metadata
        md_content += f"**Author:** {content.get('author', 'Engineering AI')}\n"
        md_content += f"**Date:** {datetime.now().strftime('%Y-%m-%d')}\n\n"
        
        # Add content
        for section in content.get("sections", []):
            md_content += f"## {section['title']}\n\n"
            md_content += f"{section['content']}\n\n"
            
            # Add subsections
            for subsection in section.get("subsections", []):
                md_content += f"### {subsection['title']}\n\n"
                md_content += f"{subsection['content']}\n\n"
                
        # Add figures
        for figure in content.get("figures", []):
            md_content += f"![{figure['caption']}]({figure['path']})\n\n"
            
        # Add tables
        for table in content.get("tables", []):
            # Add headers
            md_content += "| " + " | ".join(table["headers"]) + " |\n"
            md_content += "| " + " | ".join(["---"] * len(table["headers"])) + " |\n"
            
            # Add data
            for row in table["data"]:
                md_content += "| " + " | ".join(str(cell) for cell in row) + " |\n"
            md_content += "\n"
            
        # Sanitize the file name
        sanitized_title = content.get('title', 'report').lower().replace(' ', '_').replace('\n', '_').replace(':', '_').replace('<', '_').replace('>', '_').replace('|', '_').replace('?', '_').replace('*', '_')
        output_path = os.path.join(
            self.output_dir,
            f"{sanitized_title}.md"
        )
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(md_content)
            
        return output_path
        
    def generate_specification(self,
                             spec_data: Dict[str, Any],
                             format: str = "pdf") -> str:
        """Generate a technical specification document.
        
        Args:
            spec_data: Specification data
            format: Output format
            
        Returns:
            Path to the generated document
        """
        content = {
            "title": f"Technical Specification: {spec_data.get('name', 'Unnamed')}",
            "author": "Engineering AI",
            "sections": [
                {
                    "title": "Overview",
                    "content": spec_data.get("overview", "")
                },
                {
                    "title": "Requirements",
                    "content": self._format_requirements(spec_data.get("requirements", []))
                },
                {
                    "title": "Design",
                    "content": self._format_design(spec_data.get("design", {}))
                },
                {
                    "title": "Implementation",
                    "content": self._format_implementation(spec_data.get("implementation", {}))
                },
                {
                    "title": "Testing",
                    "content": self._format_testing(spec_data.get("testing", {}))
                }
            ]
        }
        
        return self.generate_report(content, format)
        
    def _format_requirements(self, requirements: List[Dict[str, Any]]) -> str:
        """Format requirements for documentation.
        
        Args:
            requirements: List of requirements
            
        Returns:
            Formatted requirements text
        """
        text = ""
        for req in requirements:
            text += f"### {req['id']}: {req['title']}\n\n"
            text += f"{req['description']}\n\n"
            if "criteria" in req:
                text += "**Acceptance Criteria:**\n"
                for criterion in req["criteria"]:
                    text += f"- {criterion}\n"
                text += "\n"
        return text
        
    def _format_design(self, design: Dict[str, Any]) -> str:
        """Format design information for documentation.
        
        Args:
            design: Design information
            
        Returns:
            Formatted design text
        """
        text = ""
        
        # Architecture
        if "architecture" in design:
            text += "### Architecture\n\n"
            text += f"{design['architecture']}\n\n"
            
        # Components
        if "components" in design:
            text += "### Components\n\n"
            for component in design["components"]:
                text += f"#### {component['name']}\n\n"
                text += f"{component['description']}\n\n"
                
        # Interfaces
        if "interfaces" in design:
            text += "### Interfaces\n\n"
            for interface in design["interfaces"]:
                text += f"#### {interface['name']}\n\n"
                text += f"{interface['description']}\n\n"
                
        return text
        
    def _format_implementation(self, implementation: Dict[str, Any]) -> str:
        """Format implementation details for documentation.
        
        Args:
            implementation: Implementation details
            
        Returns:
            Formatted implementation text
        """
        text = ""
        
        # Code structure
        if "code_structure" in implementation:
            text += "### Code Structure\n\n"
            text += f"{implementation['code_structure']}\n\n"
            
        # Dependencies
        if "dependencies" in implementation:
            text += "### Dependencies\n\n"
            for dep in implementation["dependencies"]:
                text += f"- {dep['name']}: {dep['version']}\n"
            text += "\n"
            
        # Configuration
        if "configuration" in implementation:
            text += "### Configuration\n\n"
            text += f"{implementation['configuration']}\n\n"
            
        return text
        
    def _format_testing(self, testing: Dict[str, Any]) -> str:
        """Format testing information for documentation.
        
        Args:
            testing: Testing information
            
        Returns:
            Formatted testing text
        """
        text = ""
        
        # Test strategy
        if "strategy" in testing:
            text += "### Test Strategy\n\n"
            text += f"{testing['strategy']}\n\n"
            
        # Test cases
        if "test_cases" in testing:
            text += "### Test Cases\n\n"
            for case in testing["test_cases"]:
                text += f"#### {case['id']}: {case['name']}\n\n"
                text += f"**Description:** {case['description']}\n\n"
                text += "**Steps:**\n"
                for step in case["steps"]:
                    text += f"1. {step}\n"
                text += "\n"
                text += f"**Expected Result:** {case['expected_result']}\n\n"
                
        return text 